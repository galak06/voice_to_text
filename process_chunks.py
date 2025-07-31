#!/usr/bin/env python3
"""
Single-Process Chunk Processor for Hebrew Audio Transcription
Processes audio chunks one by one for maximum stability
"""

import argparse
import logging
import os
import sys
from pathlib import Path
from typing import List, Dict
from dotenv import load_dotenv

from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TimeElapsedColumn
from rich.table import Table
from rich.panel import Panel

# Import our transcription module
from transcribe_hebrew import HebrewWordSpeakerTranscriber

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize Rich console for better output
console = Console()


def process_single_chunk(chunk_path: Path, model_size: str, output_dir: Path, 
                        enable_cache: bool, enable_optimization: bool, 
                        enable_text_improvements: bool) -> Dict:
    """
    Process a single audio chunk.
    
    Args:
        chunk_path: Path to the audio chunk
        model_size: Whisper model size
        output_dir: Output directory for results
        enable_cache: Enable model caching
        enable_optimization: Enable performance optimizations
        enable_text_improvements: Enable text improvements
        
    Returns:
        Dictionary with processing results
    """
    try:
        # Import utilities for proper resource management
        try:
            from utils import cleanup_resources
        except ImportError as e:
            console.print(f"[red]❌ Failed to import utils module: {e}[/red]")
            console.print(f"[red]   Make sure utils.py exists in the current directory[/red]")
            return {
                "chunk": chunk_path.name,
                "output": None,
                "status": "error",
                "error": f"Import error: {e}"
            }
        
        # Clean up memory before processing each chunk
        cleanup_resources()
        
        # Store original thread count to restore later
        import torch
        original_threads = torch.get_num_threads()
        
        # Limit torch threads to 1 for stability
        torch.set_num_threads(1)
        
        # Create output directory
        output_dir.mkdir(exist_ok=True)
        
        # Generate output filename
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = output_dir / f"{chunk_path.stem}_transcribed_{timestamp}.docx"
        
        console.print(f"[blue]🎵 Processing: {chunk_path.name}[/blue]")
        console.print(f"[blue]🧹 Memory cleaned, starting fresh...[/blue]")
        
        # Create transcriber (model will be loaded in transcribe_file)
        transcriber = HebrewWordSpeakerTranscriber(
            model_size=model_size,
            language="he",
            use_fp16=True,
            num_workers=1,  # Use 1 worker per chunk to avoid conflicts
            enable_cache=enable_cache,
            enable_optimization=enable_optimization,
            enable_text_improvements=enable_text_improvements
        )
        
        # Transcribe chunk
        transcriber.transcribe_file(str(chunk_path), str(output_file))
        
        # Clean up transcriber object and restore thread count
        del transcriber
        torch.set_num_threads(original_threads)
        cleanup_resources()
        
        console.print(f"[green]✅ {chunk_path.name} completed successfully![/green]")
        
        return {
            "chunk": chunk_path.name,
            "output": str(output_file),
            "status": "success",
            "error": None
        }
        
    except KeyboardInterrupt:
        console.print(f"[yellow]⚠️  Processing cancelled for {chunk_path.name}[/yellow]")
        cleanup_resources()
        raise
    except FileNotFoundError as e:
        console.print(f"[red]❌ File not found: {chunk_path.name} - {e}[/red]")
        cleanup_resources()
        return {
            "chunk": chunk_path.name,
            "output": None,
            "status": "error",
            "error": f"File not found: {e}"
        }
    except PermissionError as e:
        console.print(f"[red]❌ Permission denied: {chunk_path.name} - {e}[/red]")
        cleanup_resources()
        return {
            "chunk": chunk_path.name,
            "output": None,
            "status": "error",
            "error": f"Permission denied: {e}"
        }
    except Exception as e:
        console.print(f"[red]❌ Unexpected error processing {chunk_path.name}: {e}[/red]")
        cleanup_resources()
        return {
            "chunk": chunk_path.name,
            "output": None,
            "status": "error",
            "error": str(e)
        }


class ChunkProcessor:
    """
    Processes audio chunks one by one for maximum stability.
    """
    
    def __init__(self, model_size: str = "medium", 
                 enable_cache: bool = True, enable_optimization: bool = True,
                 enable_text_improvements: bool = True, enable_checkpoints: bool = True):
        """
        Initialize the chunk processor.
        
        Args:
            model_size: Whisper model size
            enable_cache: Enable model caching
            enable_optimization: Enable performance optimizations
            enable_text_improvements: Enable text improvements
        """
        self.model_size = model_size
        self.enable_cache = enable_cache
        self.enable_optimization = enable_optimization
        self.enable_text_improvements = enable_text_improvements
        self.enable_checkpoints = enable_checkpoints
        self.results = {}
        
        # Initialize checkpoint manager if enabled
        self.checkpoint_manager = None
        if self.enable_checkpoints:
            try:
                from checkpoint_manager import CheckpointManager
                self.checkpoint_manager = CheckpointManager()
            except ImportError:
                console.print("[yellow]⚠️  Checkpoint manager not available, continuing without checkpoints[/yellow]")
    
    def process_chunks_sequential(self, chunk_dir: Path, output_dir: Path = Path("output/chunks")) -> Dict:
        """
        Process chunks one by one for maximum stability.
        
        Args:
            chunk_dir: Directory containing audio chunks
            output_dir: Output directory for results
            
        Returns:
            Dictionary with processing results
        """
        # Find all audio chunks
        audio_extensions = {'.mp3', '.wav', '.m4a', '.flac', '.ogg'}
        chunk_files = [f for f in chunk_dir.iterdir() 
                      if f.is_file() and f.suffix.lower() in audio_extensions]
        
        if not chunk_files:
            console.print(f"[red]❌ No audio chunks found in {chunk_dir}[/red]")
            return {}
        
        # Import natural sorting utility
        from utils import natural_sort_files
        
        # Sort chunks by name using natural sorting
        chunk_files = natural_sort_files(chunk_files)
        
        # Initialize checkpoint session if available
        if self.checkpoint_manager:
            settings = {
                "model_size": self.model_size,
                "enable_cache": self.enable_cache,
                "enable_optimization": self.enable_optimization,
                "enable_text_improvements": self.enable_text_improvements
            }
            self.checkpoint_manager.initialize_session(len(chunk_files), settings)
            
            # Check if we can resume from previous session
            if self.checkpoint_manager.can_resume():
                remaining_chunks = self.checkpoint_manager.get_remaining_chunks([f.name for f in chunk_files])
                if remaining_chunks:
                    console.print(f"[blue]🔄 Resuming from checkpoint: {len(remaining_chunks)} chunks remaining[/blue]")
                    # Filter chunk_files to only include remaining chunks
                    chunk_files = [f for f in chunk_files if f.name in remaining_chunks]
        
        console.print(f"[blue]Found {len(chunk_files)} audio chunks to process[/blue]")
        console.print(f"[blue]Processing one by one with {self.model_size} model[/blue]")
        console.print(f"[blue]🔄 Automatic continuation enabled - will process all chunks[/blue]")
        
        # Create output directory
        output_dir.mkdir(exist_ok=True)
        
        results = {}
        
        # Process chunks one by one
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
            TimeElapsedColumn(),
            console=console
        ) as progress:
            task = progress.add_task("Processing chunks one by one...", total=len(chunk_files))
            
            for i, chunk_file in enumerate(chunk_files, 1):
                console.print(f"\n[blue]📁 Processing chunk {i}/{len(chunk_files)}: {chunk_file.name}[/blue]")
                console.print(f"[blue]⏱️  Estimated time for this chunk: 10-15 minutes[/blue]")
                
                # Mark chunk as started in checkpoint manager
                if self.checkpoint_manager:
                    self.checkpoint_manager.start_chunk(chunk_file.name)
                
                # Process single chunk
                result = process_single_chunk(
                    chunk_file, self.model_size, output_dir,
                    self.enable_cache, self.enable_optimization, self.enable_text_improvements
                )
                
                results[chunk_file.name] = result
                progress.update(task, advance=1)
                
                # Update checkpoint manager
                if self.checkpoint_manager:
                    output_file = result.get("output", "")
                    status = "success" if result["status"] == "success" else "failed"
                    self.checkpoint_manager.complete_chunk(chunk_file.name, output_file, status)
                
                if result["status"] == "success":
                    console.print(f"[green]✅ {chunk_file.name} - Completed successfully[/green]")
                    console.print(f"[blue]📄 Output saved: {result['output']}[/blue]")
                else:
                    console.print(f"[red]❌ {chunk_file.name} - Error: {result['error']}[/red]")
                
                # Show progress summary
                completed = sum(1 for r in results.values() if r["status"] == "success")
                failed = sum(1 for r in results.values() if r["status"] == "error")
                remaining = len(chunk_files) - i
                
                console.print(f"[blue]📊 Progress: {completed} completed, {failed} failed, {remaining} remaining[/blue]")
                
                # Add a small delay between chunks to ensure clean state
                if i < len(chunk_files):  # Don't delay after the last chunk
                    console.print(f"[blue]⏳ Preparing for next chunk...[/blue]")
                    import time
                    time.sleep(2)
        
        return results
    
    def merge_results(self, results: Dict, output_file: str = "merged_transcription.docx"):
        """
        Merge all chunk results into a single document.
        
        Args:
            results: Results from chunk processing
            output_file: Output merged file path
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from datetime import datetime
            
            # Create merged document
            merged_doc = Document()
            
            # Add title
            title = merged_doc.add_heading("Merged Hebrew Transcription", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add timestamp
            timestamp = merged_doc.add_paragraph(f"Merged: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            timestamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add separator
            merged_doc.add_paragraph("_" * 60)
            
            # Process successful results in order
            successful_results = [
                (chunk_name, result) for chunk_name, result in results.items()
                if result["status"] == "success"
            ]
            
            # Sort by chunk number
            successful_results.sort(key=lambda x: x[0])
            
            for chunk_name, result in successful_results:
                if result["output"] and Path(result["output"]).exists():
                    # Add chunk header
                    chunk_header = merged_doc.add_paragraph()
                    chunk_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    chunk_header_run = chunk_header.add_run(f"Chunk: {chunk_name}")
                    chunk_header_run.bold = True
                    chunk_header_run.font.size = Pt(12)
                    chunk_header_run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
                    
                    # Read and add content from chunk document
                    chunk_doc = Document(result["output"])
                    
                    # Extract paragraphs from chunk document
                    for para in chunk_doc.paragraphs:
                        if para.text.strip() and not para.text.startswith("Hebrew Transcription:"):
                            new_para = merged_doc.add_paragraph()
                            new_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            new_para.add_run(para.text)
                    
                    # Add separator between chunks
                    merged_doc.add_paragraph()
            
            # Set RTL for Hebrew
            for paragraph in merged_doc.paragraphs:
                try:
                    paragraph._element.get_or_add_pPr().get_or_add_bidi().val = 1
                except AttributeError:
                    # Fallback: try alternative method for RTL
                    try:
                        paragraph._element.get_or_add_pPr().get_or_add_textDirection().val = "rtl"
                    except AttributeError:
                        # If RTL setting fails, continue without it
                        pass
            
            # Save merged document
            merged_doc.save(output_file)
            console.print(f"[green]✅ Merged transcription saved: {output_file}[/green]")
            
        except Exception as e:
            console.print(f"[red]❌ Error merging results: {e}[/red]")


def main():
    """Main function to handle command line arguments and run chunk processing."""
    parser = argparse.ArgumentParser(
        description="Process audio chunks one by one for maximum stability",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s voice_chunks
  %(prog)s voice_chunks --model medium
  %(prog)s voice_chunks --model tiny --no-merge
        """
    )
    
    parser.add_argument("chunk_dir", help="Directory containing audio chunks")
    parser.add_argument("--model", choices=["tiny", "base", "small", "medium", "large"], 
                       default="medium", help="Whisper model size (default: medium)")
    parser.add_argument("--output-dir", default="output/chunks",
                       help="Output directory for results (default: output/chunks)")
    parser.add_argument("--no-cache", action="store_true", help="Disable model caching")
    parser.add_argument("--no-optimization", action="store_true", help="Disable performance optimizations")
    parser.add_argument("--no-text-improvements", action="store_true", 
                       help="Disable text improvements")
    parser.add_argument("--no-merge", action="store_true", 
                       help="Don't merge results into single document")
    parser.add_argument("--merge-output", default="merged_transcription.docx",
                       help="Output filename for merged document (default: merged_transcription.docx)")
    
    args = parser.parse_args()
    
    # Import validation utility
    from utils import validate_directory
    
    # Validate chunk directory
    chunk_dir_path = Path(args.chunk_dir)
    if not validate_directory(chunk_dir_path):
        console.print(f"[red]❌ Invalid chunk directory: {args.chunk_dir}[/red]")
        console.print(f"[red]   Directory must exist and be accessible[/red]")
        sys.exit(1)
    
    # Set optimization flags
    enable_cache = not args.no_cache
    enable_optimization = not args.no_optimization
    enable_text_improvements = not args.no_text_improvements
    
    # Create processor
    processor = ChunkProcessor(
        model_size=args.model,
        enable_cache=enable_cache,
        enable_optimization=enable_optimization,
        enable_text_improvements=enable_text_improvements
    )
    
    # Process chunks
    console.print(f"🎵 Processing chunks one by one with {args.model} model...")
    results = processor.process_chunks_sequential(Path(args.chunk_dir), Path(args.output_dir))
    
    # Show results summary
    successful = sum(1 for r in results.values() if r["status"] == "success")
    failed = len(results) - successful
    
    console.print(f"\n[green]📊 Processing Summary:[/green]")
    console.print(f"   ✅ Successful: {successful}")
    console.print(f"   ❌ Failed: {failed}")
    console.print(f"   📁 Output directory: {args.output_dir}")
    
    # Merge results if requested
    if not args.no_merge and successful > 0:
        console.print(f"\n[blue]🔗 Merging results...[/blue]")
        processor.merge_results(results, args.merge_output)
    
    console.print(f"\n[green]🎯 Processing complete![/green]")


if __name__ == "__main__":
    from datetime import datetime
    main() 