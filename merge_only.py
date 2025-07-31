#!/usr/bin/env python3
"""
Simple script to merge existing chunk transcription files into one document.
"""

import argparse
import logging
from pathlib import Path
from typing import Dict
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


def merge_existing_chunks(chunk_dir: Path, output_file: str = "merged_transcription.docx"):
    """
    Merge existing chunk transcription files into a single document.
    
    Args:
        chunk_dir: Directory containing chunk transcription files
        output_file: Output merged file path
    """
    try:
        # Find all .docx files in the chunk directory
        docx_files = list(chunk_dir.glob("*.docx"))
        
        # Filter out temporary files (starting with ~$)
        docx_files = [f for f in docx_files if not f.name.startswith("~$")]
        
        if not docx_files:
            print(f"❌ No .docx files found in {chunk_dir}")
            return
        
        # Import natural sorting utility
        from utils import natural_sort_files
        
        # Sort files by name using natural sorting
        docx_files = natural_sort_files(docx_files)
        
        print(f"📁 Found {len(docx_files)} transcription files to merge")
        
        # Create merged document
        merged_doc = Document()
        
        # Add title
        title = merged_doc.add_heading("Merged Hebrew Transcription", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add timestamp
        timestamp = merged_doc.add_paragraph(f"Merged: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        timestamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add separator
        merged_doc.add_paragraph("_" * 60)
        
        # Process each chunk file
        for i, docx_file in enumerate(docx_files, 1):
            print(f"📄 Processing {i}/{len(docx_files)}: {docx_file.name}")
            
            try:
                # Add chunk header
                chunk_header = merged_doc.add_paragraph()
                chunk_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                chunk_header_run = chunk_header.add_run(f"Chunk: {docx_file.stem}")
                chunk_header_run.bold = True
                chunk_header_run.font.size = Pt(12)
                chunk_header_run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
                
                # Read chunk document
                chunk_doc = Document(docx_file)
                
                # Extract paragraphs from chunk document
                for para in chunk_doc.paragraphs:
                    if para.text.strip() and not para.text.startswith("Hebrew Transcription:"):
                        # Skip "Sentence X:" lines
                        if not para.text.strip().startswith("Sentence ") and not para.text.strip().endswith(":"):
                            new_para = merged_doc.add_paragraph()
                            new_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            new_para.add_run(para.text)
                
                # Add separator between chunks
                merged_doc.add_paragraph()
                
                print(f"✅ Added content from {docx_file.name}")
                
            except Exception as e:
                print(f"❌ Error processing {docx_file.name}: {e}")
        
        # Set RTL for Hebrew
        for paragraph in merged_doc.paragraphs:
            try:
                paragraph._element.get_or_add_pPr().get_or_add_bidi().val = 1
            except AttributeError:
                # Fallback: try alternative method for RTL
                try:
                    paragraph._element.get_or_add_pPr().get_or_add_textDirection().val = "rtl"
                except AttributeError:
                    # If RTL setting fails, continue without it
                    pass
        
        # Save merged document
        merged_doc.save(output_file)
        print(f"✅ Merged transcription saved: {output_file}")
        print(f"📊 Total paragraphs in merged document: {len(merged_doc.paragraphs)}")
        
    except Exception as e:
        print(f"❌ Error merging results: {e}")


def main():
    """Main function to handle command line arguments and run merging."""
    parser = argparse.ArgumentParser(
        description="Merge existing chunk transcription files into a single document",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s output_chunks
  %(prog)s output_chunks --output merged_new.docx
        """
    )
    
    parser.add_argument("chunk_dir", help="Directory containing chunk transcription files")
    parser.add_argument("--output", default="output/merged/merged_transcription.docx",
                       help="Output filename for merged document (default: output/merged/merged_transcription.docx)")
    
    args = parser.parse_args()
    
    # Import validation utility
    from utils import validate_directory
    
    # Validate chunk directory
    chunk_dir_path = Path(args.chunk_dir)
    if not validate_directory(chunk_dir_path):
        print(f"❌ Invalid chunk directory: {args.chunk_dir}")
        print(f"   Directory must exist and be accessible")
        return
    
    print(f"🔗 Merging transcription files from {args.chunk_dir}...")
    merge_existing_chunks(Path(args.chunk_dir), args.output)
    print(f"🎯 Merging complete!")


if __name__ == "__main__":
    main() 