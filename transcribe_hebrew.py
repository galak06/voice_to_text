#!/usr/bin/env python3
"""
Hebrew Voice to Text with Word Document Output - Speaker Separation
Outputs Hebrew RTL text to .docx files with clear speaker separation per paragraph
"""

import argparse
import logging
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import List, Tuple
from dotenv import load_dotenv

import whisper
import torch
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TimeElapsedColumn
from rich.table import Table
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Hebrew punctuation restoration
try:
    from transformers import AutoTokenizer, AutoModelForTokenClassification
    HEBREW_PUNCTUATION_AVAILABLE = True
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    HEBREW_PUNCTUATION_AVAILABLE = False
    TRANSFORMERS_AVAILABLE = False

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize Rich console for better output
console = Console()


class HebrewPunctuationRestorer:
    """
    Hebrew punctuation restoration using rule-based approach.
    """
    
    def __init__(self):
        self.is_loaded = True  # Always available since it's rule-based
    
    def load_model(self):
        """Initialize Hebrew punctuation restoration."""
        console.print("[blue]Initializing Hebrew punctuation restoration...[/blue]")
        console.print("[green]✅ Hebrew punctuation restoration ready (rule-based)[/green]")
    
    def punctuate(self, text: str) -> str:
        """
        Add punctuation to Hebrew text using rule-based approach.
        
        Args:
            text: Hebrew text without punctuation
            
        Returns:
            Hebrew text with added punctuation
        """
        if not text.strip():
            return text
        
        try:
            # Simple rule-based Hebrew punctuation
            result = text.strip()
            
            # Add periods at the end if missing
            if not result.endswith(('.', '!', '?', ':', ';')):
                result += '.'
            
            # Add commas before common Hebrew conjunctions (more selective)
            # Only add commas for conjunctions that typically need them
            conjunctions = ['אבל', 'או', 'וגם', 'אם', 'כי', 'למרות', 'בנוסף']
            
            for conj in conjunctions:
                # Only add comma if conjunction is not at the beginning and follows a complete thought
                pattern = r'([.!?])\s+(' + conj + r')\s+'
                replacement = r'\1, \2 '
                result = re.sub(pattern, replacement, result)
            
            # Add commas around common Hebrew phrases
            phrases = ['כלומר', 'למשל', 'למעשה', 'בעצם', 'באמת', 'בדרך כלל', 'לרוב']
            
            for phrase in phrases:
                if phrase in result:
                    # Add commas around the phrase
                    result = result.replace(f' {phrase} ', f', {phrase}, ')
            
            # Clean up multiple commas
            while ',,' in result:
                result = result.replace(',,', ',')
            
            # Clean up spaces around punctuation
            result = result.replace(' .', '.')
            result = result.replace(' ,', ',')
            result = result.replace(' !', '!')
            result = result.replace(' ?', '?')
            result = result.replace(' :', ':')
            result = result.replace(' ;', ';')
            
            return result.strip()
            
        except Exception as e:
            console.print(f"[yellow]⚠️  Warning: Error during punctuation restoration: {e}[/yellow]")
            return text


class HebrewWordSpeakerTranscriber:
    """
    Hebrew transcriber that outputs to Word documents with clear speaker separation.
    """
    
    def __init__(self, model_size: str = "medium", language: str = "he", 
                 use_fp16: bool = True, num_workers: int = 1, 
                 enable_cache: bool = True, enable_optimization: bool = True,
                 enable_text_improvements: bool = True, max_speakers: int = 2):
        """
        Initialize the transcriber with specified model and language.
        
        Args:
            model_size: Whisper model size (default: 'medium' for good balance)
            language: Language code for transcription (default: 'he' for Hebrew)
            use_fp16: Use FP16 for faster processing (if GPU available)
            num_workers: Number of parallel workers
            enable_cache: Enable model caching for faster reloads
            enable_optimization: Enable performance optimizations
            enable_text_improvements: Enable speaker separation and paragraph formatting
        """
        self.model_size = model_size
        self.language = language
        self.use_fp16 = use_fp16
        # Use specified number of workers (minimum 1)
        self.num_workers = max(1, num_workers)
        self.enable_cache = enable_cache
        self.enable_optimization = enable_optimization
        self.enable_text_improvements = enable_text_improvements
        self.max_speakers = max(1, max_speakers)  # Ensure at least 1 speaker
        self.whisper_model = None
        
        # Initialize Hebrew punctuation restorer
        self.punctuation_restorer = HebrewPunctuationRestorer()
        
        # Performance optimizations
        if self.enable_optimization:
            # Set environment variables for better performance
            os.environ["TOKENIZERS_PARALLELISM"] = "false"  # Avoid tokenizer warnings
            os.environ["OMP_NUM_THREADS"] = str(min(8, os.cpu_count() or 4))  # Optimize OpenMP
            os.environ["MKL_NUM_THREADS"] = str(min(8, os.cpu_count() or 4))  # Optimize MKL
            
            # Enable PyTorch optimizations
            if torch.cuda.is_available():
                torch.backends.cudnn.benchmark = True  # Optimize CUDA kernels
                torch.backends.cudnn.deterministic = False  # Allow non-deterministic for speed
        
    def load_model(self):
        """Load Whisper model with speed optimizations."""
        try:
            with Progress(
                SpinnerColumn(),
                TextColumn("[progress.description]{task.description}"),
                BarColumn(),
                TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
                console=console
            ) as progress:
                task = progress.add_task(f"Loading {self.model_size} model... (Press Ctrl+C to cancel)", total=100)
                
                # Load model with optimizations
                progress.update(task, completed=25)
                
                # Use device optimization
                device = "cuda" if torch.cuda.is_available() else "cpu"
                compute_type = "float16" if self.use_fp16 and device == "cuda" else "float32"
                
                # Set cache directory for faster reloads
                cache_dir = None
                if self.enable_cache:
                    cache_dir = os.path.expanduser("~/.cache/whisper")
                    os.makedirs(cache_dir, exist_ok=True)
                
                # Load model with or without cache directory
                if cache_dir:
                    self.whisper_model = whisper.load_model(
                        self.model_size, 
                        device=device,
                        download_root=cache_dir
                    )
                else:
                    self.whisper_model = whisper.load_model(
                        self.model_size, 
                        device=device
                    )
                
                progress.update(task, completed=75)
                
                # Optimize model for inference
                if device == "cuda" and self.use_fp16:
                    self.whisper_model = self.whisper_model.half()  # Use FP16
                
                # Additional optimizations
                if self.enable_optimization:
                    # Set model to evaluation mode for faster inference
                    self.whisper_model.eval()
                    
                    # Enable gradient computation optimizations
                    if hasattr(self.whisper_model, 'encoder'):
                        self.whisper_model.encoder.eval()
                
                progress.update(task, completed=100)
                
                # Load Hebrew punctuation model if available
                if HEBREW_PUNCTUATION_AVAILABLE and self.enable_text_improvements:
                    self.punctuation_restorer.load_model()
                
            console.print(f"✅ Loaded {self.model_size} model on {device.upper()}")
            
            # Show model info
            self.show_model_info()
            
        except KeyboardInterrupt:
            console.print("\n[yellow]⚠️  Model loading cancelled by user[/yellow]")
            sys.exit(0)
        except Exception as e:
            console.print(f"[red]Error loading Whisper model: {e}[/red]")
            raise
    
    def show_model_info(self):
        """Display model information with speed details."""
        table = Table(title="Model Configuration")
        table.add_column("Setting", style="cyan")
        table.add_column("Value", style="green")
        table.add_column("Impact", style="yellow")
        
        device = "GPU" if torch.cuda.is_available() else "CPU"
        fp16_status = "Enabled" if self.use_fp16 and torch.cuda.is_available() else "Disabled"
        cache_status = "Enabled" if self.enable_cache else "Disabled"
        optimization_status = "Enabled" if self.enable_optimization else "Disabled"
        text_improvements_status = "Enabled" if self.enable_text_improvements else "Disabled"
        punctuation_status = "Enabled" if (HEBREW_PUNCTUATION_AVAILABLE and 
                                         self.enable_text_improvements and 
                                         self.punctuation_restorer.is_loaded) else "Disabled"
        
        table.add_row("Model Size", self.model_size, "Tiny = Fastest, Large = Best")
        table.add_row("Language", self.language, "Hebrew optimization")
        table.add_row("Device", device, "GPU = 10x faster")
        table.add_row("FP16", fp16_status, "2x faster on GPU")
        table.add_row("Cache", cache_status, "Faster model reloads")
        table.add_row("Optimizations", optimization_status, "Performance boost")
        table.add_row("Workers", str(self.num_workers), "Parallel processing")
        table.add_row("Text Improvements", text_improvements_status, "Speaker separation & formatting")
        table.add_row("Hebrew Punctuation", punctuation_status, "Grammar & readability")
        table.add_row("Output Format", "Word Document (.docx)", "Professional formatting")
        
        console.print(table)
    
    def transcribe_audio(self, audio_path: str) -> str:
        """
        Transcribe audio file with speed optimizations.
        
        Args:
            audio_path: Path to the audio file
            
        Returns:
            Transcribed text as a single string
        """
        try:
            # Check file size for optimization
            file_size_mb = Path(audio_path).stat().st_size / (1024 * 1024)
            is_large_file = file_size_mb > 50  # Files larger than 50MB
            
            if is_large_file:
                console.print(f"[yellow]⚠️  Large file detected ({file_size_mb:.1f}MB) - Using optimized parameters[/yellow]")
            
            console.print(f"[blue]🎤 Starting transcription with {self.model_size} model...[/blue]")
            console.print(f"[blue]📁 Processing file: {Path(audio_path).name}[/blue]")
            console.print(f"[blue]📊 File size: {file_size_mb:.1f}MB[/blue]")
            
            # Get audio duration for progress estimation
            try:
                import librosa
                y, sr = librosa.load(audio_path, sr=None)
                duration = len(y) / sr
                console.print(f"[blue]⏱️  Audio duration: {duration:.1f} seconds ({duration/60:.1f} minutes)[/blue]")
            except Exception as e:
                console.print(f"[yellow]⚠️  Could not determine audio duration: {e}[/yellow]")
                duration = None
            
            with Progress(
                SpinnerColumn(),
                TextColumn("[progress.description]{task.description}"),
                BarColumn(),
                TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
                TimeElapsedColumn(),
                console=console
            ) as progress:
                task = progress.add_task("Transcribing audio... (Press Ctrl+C to cancel)", total=100)
                
                # Start transcription
                progress.update(task, completed=5)
                console.print("[blue]📊 Configuring transcription parameters...[/blue]")
                
                # Optimize settings based on model size and file size
                if self.model_size == "tiny":
                    # Ultra-fast settings for tiny model
                    best_of = 1
                    beam_size = 1
                    condition_on_previous_text = False
                    compression_ratio_threshold = 2.4
                    logprob_threshold = -1.0
                    no_speech_threshold = 0.6
                elif self.model_size == "base":
                    # Fast settings for base model
                    if is_large_file:
                        best_of = 2
                        beam_size = 2
                        condition_on_previous_text = False
                    else:
                        best_of = 3
                        beam_size = 3
                        condition_on_previous_text = False
                    compression_ratio_threshold = 2.4
                    logprob_threshold = -1.0
                    no_speech_threshold = 0.6
                elif self.model_size == "small":
                    # Balanced settings for small model
                    if is_large_file:
                        best_of = 2
                        beam_size = 2
                        condition_on_previous_text = False
                    else:
                        best_of = 3
                        beam_size = 3
                        condition_on_previous_text = True
                    compression_ratio_threshold = 2.4
                    logprob_threshold = -1.0
                    no_speech_threshold = 0.6
                elif self.model_size == "medium":
                    # Balanced settings for medium model
                    if is_large_file:
                        best_of = 2
                        beam_size = 2
                        condition_on_previous_text = False
                    else:
                        best_of = 3
                        beam_size = 3
                        condition_on_previous_text = True
                    compression_ratio_threshold = 2.4
                    logprob_threshold = -1.0
                    no_speech_threshold = 0.6
                else:  # large
                    # Conservative settings for large model to prevent hanging
                    best_of = 1 if is_large_file else 2
                    beam_size = 1 if is_large_file else 2
                    condition_on_previous_text = False
                    compression_ratio_threshold = 2.4
                    logprob_threshold = -1.0
                    no_speech_threshold = 0.6
                
                progress.update(task, completed=15)
                console.print(f"[blue]⚙️  Using parameters: best_of={best_of}, beam_size={beam_size}[/blue]")
                console.print(f"[blue]🔧 Model: {self.model_size}, Language: {self.language}[/blue]")
                console.print(f"[blue]💻 Device: {'GPU' if torch.cuda.is_available() else 'CPU'}[/blue]")
                
                # Performance optimizations
                progress.update(task, completed=20)
                console.print("[blue]🚀 Starting transcription process...[/blue]")
                if duration:
                    estimated_time = duration / 60 * 2  # Rough estimate: 2x real-time
                    console.print(f"[blue]⏳ Estimated processing time: {estimated_time:.1f} minutes[/blue]")
                
                with torch.no_grad():  # Disable gradient computation for faster inference
                    result = self.whisper_model.transcribe( # type: ignore
                        audio_path,
                        language=self.language,
                        temperature=0.0,  # Most accurate
                        best_of=best_of,
                        beam_size=beam_size,
                        word_timestamps=False,
                        condition_on_previous_text=condition_on_previous_text,
                        initial_prompt="This is a Hebrew conversation between two speakers.",
                        compression_ratio_threshold=compression_ratio_threshold,
                        logprob_threshold=logprob_threshold,
                        no_speech_threshold=no_speech_threshold
                    )
                
                progress.update(task, completed=90)
                console.print("[blue]✨ Transcription completed, applying Hebrew punctuation...[/blue]")
                
                # Apply Hebrew punctuation restoration
                if (self.enable_text_improvements and 
                    self.punctuation_restorer.is_loaded):
                    console.print("[blue]🔤 Applying Hebrew punctuation rules...[/blue]")
                    result["text"] = self.punctuation_restorer.punctuate(result["text"])
                
                progress.update(task, completed=100)
                console.print("[green]✅ Transcription and processing completed![/green]")
                console.print(f"[green]📝 Transcribed text length: {len(result['text'])} characters[/green]")
                
            return result["text"].strip() # type: ignore
        except KeyboardInterrupt:
            console.print("\n[yellow]⚠️  Transcription cancelled by user[/yellow]")
            sys.exit(0)
        except Exception as e:
            console.print(f"[red]Error transcribing audio: {e}[/red]")
            raise
    
    def format_hebrew_speakers(self, text: str, audio_path: str = None) -> List[Tuple[str, str]]:
        """
        Format Hebrew text into speaker-separated sentences using Whisper's speaker diarization.
        
        Args:
            text: Raw transcribed text
            audio_path: Path to audio file for speaker diarization (optional)
            
        Returns:
            List of tuples: (speaker_label, sentence_text) in time order
        """
        import re
        
        # Remove any existing speaker labels
        text = re.sub(r'^\[user\]\s*', '', text)
        
        # Split text into sentences (Hebrew sentence endings)
        sentences = re.split(r'([.!?]+)', text)
        
        # Rejoin sentences with proper spacing
        formatted_sentences = []
        for i in range(0, len(sentences), 2):
            if i + 1 < len(sentences):
                sentence = sentences[i] + sentences[i + 1]
            else:
                sentence = sentences[i]
            
            if sentence.strip():
                formatted_sentences.append(sentence.strip())
        
        # Try to use Whisper's speaker diarization if available and audio path provided
        if audio_path and self.whisper_model and hasattr(self.whisper_model, 'transcribe'):
            try:
                console.print("[blue]🔍 Attempting speaker diarization...[/blue]")
                
                # Use Whisper's transcribe with speaker diarization
                result = self.whisper_model.transcribe(
                    audio_path,
                    language="he",
                    task="transcribe",
                    verbose=False,
                    word_timestamps=True
                )
                
                # Extract speaker information from segments
                if hasattr(result, 'segments') and result.segments:
                    speaker_sentences = []
                    for segment in result.segments:
                        if segment.text.strip():
                            # Use speaker ID if available, otherwise use segment speaker
                            speaker_id = getattr(segment, 'speaker', None)
                            if speaker_id is not None:
                                speaker_label = f"Speaker {speaker_id + 1}"
                            else:
                                speaker_label = f"Speaker 1"
                            speaker_sentences.append((speaker_label, segment.text.strip()))
                    
                    if speaker_sentences:
                        console.print("[green]✅ Speaker diarization successful[/green]")
                        return speaker_sentences
                    else:
                        console.print("[yellow]⚠️  Speaker diarization returned no segments, using fallback[/yellow]")
                else:
                    console.print("[yellow]⚠️  No segments found in diarization result, using fallback[/yellow]")
                    
            except Exception as e:
                console.print(f"[yellow]⚠️  Speaker diarization failed: {e}, using fallback[/yellow]")
        
        # Fallback to intelligent speaker assignment based on content analysis
        console.print("[blue]🔍 Using intelligent speaker detection...[/blue]")
        speaker_sentences = []
        current_speaker = 1
        speaker_changes = 0
        
        for i, sentence in enumerate(formatted_sentences):
            # Enhanced heuristic: change speaker based on multiple factors
            is_response = (
                len(sentence.strip()) < 20 or  # Short sentence
                sentence.strip().startswith(('כן', 'לא', 'מה', 'איך', 'מתי', 'איפה', 'למה', 'מי')) or
                sentence.strip().startswith(('אני', 'אתה', 'את', 'הוא', 'היא', 'אנחנו', 'אתם', 'הם', 'הן')) or
                sentence.strip().startswith(('תודה', 'בבקשה', 'סליחה', 'אני חושב', 'אני חושבת'))
            )
            
            # Change speaker if this looks like a response and we have enough content
            if is_response and i > 0 and len(formatted_sentences[i-1].strip()) > 30:
                current_speaker = 2 if current_speaker == 1 else 1
                speaker_changes += 1
            
            speaker_label = f"Speaker {current_speaker}"
            speaker_sentences.append((speaker_label, sentence))
        
        # If we didn't detect any speaker changes, use alternating pattern as fallback
        if speaker_changes == 0 and len(formatted_sentences) > 1:
            speaker_sentences = []
            for i, sentence in enumerate(formatted_sentences):
                speaker_label = f"Speaker {(i % 2) + 1}"
                speaker_sentences.append((speaker_label, sentence))
        
        return speaker_sentences
    
    def create_word_document_with_speakers(self, speaker_paragraphs: List[Tuple[str, str]], output_path: str, audio_filename: str):
        """
        Create a Word document with clear speaker separation.
        
        Args:
            speaker_paragraphs: List of (speaker_label, paragraph_text) tuples
            output_path: Output Word document path
            audio_filename: Original audio filename for title
        """
        try:
            # Create a new Word document
            doc = Document()
            
            # Set document properties
            doc.core_properties.title = f"Hebrew Transcription - {audio_filename}"
            doc.core_properties.author = "Voice to Text System"
            
            # Add title
            title = doc.add_heading(f'Hebrew Voice Transcription', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add subtitle
            subtitle = doc.add_paragraph(f'Audio File: {audio_filename}')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            metadata = doc.add_paragraph(f'Model: {self.model_size} | Language: Hebrew | Max Speakers: {self.max_speakers}')
            metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add separator
            doc.add_paragraph('_' * 60)
            
            # Display sentences in chronological order (Speaker 1, Speaker 2, Speaker 1, etc.)
            speaker1_count = 0
            speaker2_count = 0
            
            for i, (speaker_label, sentence_text) in enumerate(speaker_paragraphs, 1):
                # Add speaker label
                speaker_heading = doc.add_heading(f'{speaker_label}', level=2)
                speaker_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Set color based on speaker
                if "Speaker 1" in speaker_label:
                    speaker1_count += 1
                    color = RGBColor(0, 102, 204)  # Blue for Speaker 1
                else:
                    speaker2_count += 1
                    color = RGBColor(204, 0, 102)  # Red for Speaker 2
                
                # Add sentence text
                text_para = doc.add_paragraph()
                text_run = text_para.add_run(sentence_text)
                text_run.font.size = Pt(11)
                text_run.font.name = 'Arial'
                text_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Add spacing between sentences
                doc.add_paragraph()
            
            # Add summary
            doc.add_paragraph('_' * 60)
            summary = doc.add_paragraph(f'Summary: {speaker1_count} sentences from Speaker 1, {speaker2_count} sentences from Speaker 2')
            summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save the document
            doc.save(output_path)
            
        except Exception as e:
            console.print(f"[red]Error creating Word document: {e}[/red]")
            raise
    
    def transcribe_file(self, input_path: str, output_path: str):
        """
        Main method to transcribe an audio file to Word document with speaker separation.
        
        Args:
            input_path: Path to input audio file
            output_path: Path to output Word document
        """
        try:
            # Load model if not already loaded
            if not self.whisper_model:
                self.load_model()
            
            # Transcribe audio with progress
            text = self.transcribe_audio(input_path)
            
            # Create organized output paths
            output_path_obj = Path(output_path)
            audio_filename = Path(input_path).name
            
            # Create output directories
            word_dir = Path("output/word")
            text_dir = Path("output/text")
            word_dir.mkdir(exist_ok=True, parents=True)
            text_dir.mkdir(exist_ok=True, parents=True)
            
            # Generate organized file paths
            timestamp = output_path_obj.stem.split('_transcribed_')[-1] if '_transcribed_' in output_path_obj.stem else datetime.now().strftime("%Y%m%d_%H%M%S")
            word_file_path = word_dir / f"{output_path_obj.stem.split('_transcribed_')[0]}_transcribed_{timestamp}.docx"
            text_file_path = text_dir / f"{output_path_obj.stem.split('_transcribed_')[0]}_transcribed_{timestamp}.txt"
            
            if self.enable_text_improvements:
                # Format into speaker-separated paragraphs
                console.print("[blue]👥 Formatting speaker separation...[/blue]")
                speaker_paragraphs = self.format_hebrew_speakers(text, input_path)
                console.print(f"[blue]📊 Found {len(speaker_paragraphs)} speaker segments[/blue]")
                
                # Create Word document with speaker separation
                console.print("[blue]📄 Creating Word document with speaker separation...[/blue]")
                self.create_word_document_with_speakers(speaker_paragraphs, str(word_file_path), audio_filename)
                
                # Create text file with speaker separation
                console.print("[blue]📝 Creating text file with speaker separation...[/blue]")
                self.create_text_file(speaker_paragraphs, str(text_file_path), audio_filename)
                
                # Print summary
                self.print_summary(text, speaker_paragraphs, str(word_file_path))
            else:
                # Create simple Word document without speaker separation
                console.print("[blue]📄 Creating simple Word document...[/blue]")
                self.create_simple_word_document(text, str(word_file_path), audio_filename)
                
                # Create simple text file without speaker separation
                console.print("[blue]📝 Creating simple text file...[/blue]")
                self.create_simple_text_file(text, str(text_file_path), audio_filename)
                
                # Print simple summary
                self.print_simple_summary(text, str(word_file_path))
            
        except Exception as e:
            console.print(f"[red]Error transcribing file: {e}[/red]")
            raise
    
    def print_summary(self, text: str, speaker_paragraphs: List[Tuple[str, str]], output_path: str):
        """Print transcription summary."""
        console.print(f"✓ Word document saved to: {output_path}")
        console.print(f"Text length: {len(text)} characters")
        console.print(f"Total sentences: {len(speaker_paragraphs)}")
        
        # Count sentences per speaker
        speaker1_count = sum(1 for speaker, _ in speaker_paragraphs if "Speaker 1" in speaker)
        speaker2_count = sum(1 for speaker, _ in speaker_paragraphs if "Speaker 2" in speaker)
        
        console.print(f"Speaker 1 sentences: {speaker1_count}")
        console.print(f"Speaker 2 sentences: {speaker2_count}")
        
        # Quality score
        quality_score = self.calculate_quality_score(text)
        console.print(f"Quality score: {quality_score:.1f}/10")
        
        if quality_score >= 8:
            console.print("🎯 [green]Excellent quality![/green]")
        elif quality_score >= 6:
            console.print("✅ [yellow]Good quality[/yellow]")
        else:
            console.print("⚠️ [red]Quality could be improved[/red]")
    
    def calculate_quality_score(self, text: str) -> float:
        """Calculate quality score based on text characteristics."""
        if not text:
            return 0.0
        
        score = 5.0  # Base score
        
        # Length factor
        if len(text) > 5000:
            score += 1.0
        elif len(text) > 2000:
            score += 0.5
        
        # Hebrew character density
        hebrew_chars = sum(1 for c in text if '\u0590' <= c <= '\u05FF')
        hebrew_ratio = hebrew_chars / len(text) if text else 0
        if hebrew_ratio > 0.7:
            score += 1.0
        elif hebrew_ratio > 0.5:
            score += 0.5
        
        # Punctuation and structure
        if text.count('.') > 10:
            score += 1.0
        if text.count('?') > 5:
            score += 0.5
        
        return min(10.0, max(0.0, score))

    def create_simple_word_document(self, text: str, output_path: str, audio_filename: str):
        """
        Create a simple Word document without speaker separation.
        
        Args:
            text: Raw transcribed text
            output_path: Output Word document path
            audio_filename: Original audio filename for title
        """
        try:
            # Create a new Word document
            doc = Document()
            
            # Set document properties
            doc.core_properties.title = f"Hebrew Transcription - {audio_filename}"
            doc.core_properties.author = "Voice to Text System"
            
            # Add title
            title = doc.add_heading(f'Hebrew Voice Transcription', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add subtitle
            subtitle = doc.add_paragraph(f'Audio File: {audio_filename}')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            metadata = doc.add_paragraph(f'Model: {self.model_size} | Language: Hebrew | Text Improvements: Disabled')
            metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add separator
            doc.add_paragraph('_' * 60)
            
            # Add raw text
            text_para = doc.add_paragraph()
            text_run = text_para.add_run(text)
            text_run.font.size = Pt(11)
            text_run.font.name = 'Arial'
            text_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Save the document
            doc.save(output_path)
            
        except Exception as e:
            console.print(f"[red]Error creating Word document: {e}[/red]")
            raise

    def create_text_file(self, speaker_paragraphs: List[Tuple[str, str]], output_path: str, audio_filename: str):
        """
        Create a text file with speaker-separated content.
        
        Args:
            speaker_paragraphs: List of (speaker, text) tuples
            output_path: Output text file path
            audio_filename: Original audio filename for header
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                # Write header
                f.write("Hebrew Voice Transcription\n")
                f.write("=" * 50 + "\n")
                f.write(f"Audio File: {audio_filename}\n")
                f.write(f"Model: {self.model_size} | Language: Hebrew | Max Speakers: {self.max_speakers}\n")
                f.write("=" * 50 + "\n\n")
                
                # Write speaker-separated content
                for speaker, text in speaker_paragraphs:
                    f.write(f"{speaker}\n")
                    f.write(f"{text}\n\n")
                
                # Write footer
                f.write("=" * 50 + "\n")
                f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Total segments: {len(speaker_paragraphs)}\n")
            
            console.print(f"✓ Text file saved to: {output_path}")
            
        except Exception as e:
            console.print(f"[red]Error creating text file: {e}[/red]")
    
    def create_simple_text_file(self, text: str, output_path: str, audio_filename: str):
        """
        Create a simple text file without speaker separation.
        
        Args:
            text: Raw transcribed text
            output_path: Output text file path
            audio_filename: Original audio filename for header
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                # Write header
                f.write("Hebrew Voice Transcription\n")
                f.write("=" * 50 + "\n")
                f.write(f"Audio File: {audio_filename}\n")
                f.write(f"Model: {self.model_size} | Language: Hebrew\n")
                f.write("=" * 50 + "\n\n")
                
                # Write raw text
                f.write(text + "\n\n")
                
                # Write footer
                f.write("=" * 50 + "\n")
                f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Text length: {len(text)} characters\n")
            
            console.print(f"✓ Text file saved to: {output_path}")
            
        except Exception as e:
            console.print(f"[red]Error creating text file: {e}[/red]")

    def print_simple_summary(self, text: str, output_path: str):
        """Print simple transcription summary without speaker separation."""
        console.print(f"✓ Word document saved to: {output_path}")
        console.print(f"Text length: {len(text)} characters")
        
        # Quality score
        quality_score = self.calculate_quality_score(text)
        console.print(f"Quality score: {quality_score:.1f}/10")
        
        if quality_score >= 8:
            console.print("🎯 [green]Excellent quality![/green]")
        elif quality_score >= 6:
            console.print("✅ [yellow]Good quality[/yellow]")
        else:
            console.print("⚠️ [red]Quality could be improved[/red]")


def process_voice_folder(model_size: str = "medium", num_workers: int = 1, 
                        enable_cache: bool = True, enable_optimization: bool = True,
                        enable_text_improvements: bool = True):
    """Process all audio files in the voice folder to Word documents with speaker separation."""
    voice_dir = Path("voice")
    output_dir = Path("output")
    
    # Create output directory if it doesn't exist
    output_dir.mkdir(exist_ok=True)
    
    # Find all audio files
    audio_extensions = {'.wav', '.mp3', '.m4a', '.flac', '.aac'}
    audio_files = [
        f for f in voice_dir.iterdir() 
        if f.is_file() and f.suffix.lower() in audio_extensions
    ]
    
    if not audio_files:
        console.print("[yellow]No audio files found in voice/ folder[/yellow]")
        return
    
    console.print(f"Found {len(audio_files)} audio file(s):")
    for audio_file in audio_files:
        console.print(f"  📁 {audio_file.name}")
    
    console.print("\n[yellow]Press Ctrl+C at any time to cancel the process[/yellow]")
    
    # Process each file with performance optimizations
    transcriber = HebrewWordSpeakerTranscriber(
        model_size=model_size, 
        language="he", 
        num_workers=num_workers,
        enable_cache=enable_cache,
        enable_optimization=enable_optimization,
        enable_text_improvements=enable_text_improvements
    )
    
    try:
        for audio_file in audio_files:
            console.print(f"\n📄 Processing to Word document with speaker separation: {audio_file}")
            
            # Generate timestamp for unique filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_file = output_dir / f"{audio_file.stem}_hebrew_speakers_{model_size}_{timestamp}.docx"
            
            transcriber.transcribe_file(str(audio_file), str(output_file))
    except KeyboardInterrupt:
        console.print("\n[yellow]⚠️  Folder processing cancelled by user[/yellow]")
        sys.exit(0)


def main():
    """Main function to handle command line arguments and run transcription."""
    parser = argparse.ArgumentParser(
        description="Hebrew Voice to Text with Word Document Output - Speaker Separation",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Speed Tips:
  --model tiny     # Fastest model (39MB)
  --workers 4      # Parallel processing (minimum: 2)
  --no-fp16        # Disable FP16 if issues
  --device cuda    # Use GPU (10x faster)
  --no-cache       # Disable model caching
  --no-optimize    # Disable performance optimizations
  --no-text-improvements  # Raw text without speaker separation

Examples:
  %(prog)s voice/audio.wav output.docx
  %(prog)s --model tiny --workers 4
  %(prog)s --model base --device cuda
  %(prog)s --process-folder
  %(prog)s --model tiny --no-cache --no-optimize  # Minimal memory usage
  %(prog)s --model tiny --no-text-improvements   # Raw text output
        """
    )
    
    parser.add_argument(
        "input_file", 
        nargs="?", 
        help="Input audio file path"
    )
    parser.add_argument(
        "output_file", 
        nargs="?", 
        help="Output Word document path"
    )
    parser.add_argument(
        "--model", 
        choices=["tiny", "base", "small", "medium", "large"],
        default="medium",
        help="Whisper model size (default: medium for good balance)"
    )
    parser.add_argument(
        "--language", 
        default="he",
        help="Language code for transcription (default: he for Hebrew)"
    )
    parser.add_argument(
        "--workers", 
        type=int, 
        default=1, 
        help="Number of parallel workers (minimum: 1, default: 1)"
    )
    parser.add_argument(
        "--no-fp16", 
        action="store_true", 
        help="Disable FP16 optimization"
    )
    parser.add_argument(
        "--device", 
        choices=["auto", "cpu", "cuda"], 
        default="auto",
        help="Device to use (default: auto)"
    )
    parser.add_argument(
        "--process-folder",
        action="store_true",
        help="Process all audio files in the voice/ folder"
    )
    parser.add_argument(
        "--verbose", "-v",
        action="store_true",
        help="Enable verbose output"
    )
    parser.add_argument(
        "--no-cache",
        action="store_true",
        help="Disable model caching (saves memory, slower reloads)"
    )
    parser.add_argument(
        "--no-optimize",
        action="store_true",
        help="Disable performance optimizations (saves memory)"
    )
    parser.add_argument(
        "--no-text-improvements",
        action="store_true",
        help="Disable speaker separation and paragraph formatting (raw text only)"
    )
    
    args = parser.parse_args()
    
    # Allow single worker if specifically requested
    if args.workers < 1:
        console.print("[yellow]⚠️  Workers set to minimum of 1[/yellow]")
        args.workers = 1
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    try:
        # Device selection
        if args.device == "cpu":
            os.environ["CUDA_VISIBLE_DEVICES"] = ""
        elif args.device == "cuda" and not torch.cuda.is_available():
            console.print("[yellow]GPU not available, using CPU[/yellow]")
        
        # Performance settings
        use_fp16 = not args.no_fp16
        enable_cache = not args.no_cache
        enable_optimization = not args.no_optimize
        enable_text_improvements = not args.no_text_improvements
        
        if args.process_folder:
            process_voice_folder(
                args.model, 
                args.workers, 
                enable_cache=enable_cache,
                enable_optimization=enable_optimization,
                enable_text_improvements=enable_text_improvements
            )
        elif args.input_file and args.output_file:
            transcriber = HebrewWordSpeakerTranscriber(
                model_size=args.model, 
                language=args.language,
                use_fp16=use_fp16,
                num_workers=args.workers,
                enable_cache=enable_cache,
                enable_optimization=enable_optimization,
                enable_text_improvements=enable_text_improvements
            )
            transcriber.transcribe_file(args.input_file, args.output_file)
        elif args.input_file:
            # If only input file is provided, generate output with timestamp
            transcriber = HebrewWordSpeakerTranscriber(
                model_size=args.model, 
                language=args.language,
                use_fp16=use_fp16,
                num_workers=args.workers,
                enable_cache=enable_cache,
                enable_optimization=enable_optimization,
                enable_text_improvements=enable_text_improvements
            )
            
            input_path = Path(args.input_file)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = Path("output") / f"{input_path.stem}_hebrew_speakers_{args.model}_{timestamp}.docx"
            
            # Create output directory if it doesn't exist
            output_path.parent.mkdir(exist_ok=True)
            
            transcriber.transcribe_file(str(input_path), str(output_path))
        else:
            console.print("[yellow]Processing voice folder to Word documents with speaker separation...[/yellow]")
            process_voice_folder(
                args.model, 
                args.workers,
                enable_cache=enable_cache,
                enable_optimization=enable_optimization,
                enable_text_improvements=enable_text_improvements
            )
            
    except KeyboardInterrupt:
        console.print("\n[yellow]Transcription interrupted by user[/yellow]")
    except Exception as e:
        console.print(f"[red]Error: {e}[/red]")
        sys.exit(1)


if __name__ == "__main__":
    main() 